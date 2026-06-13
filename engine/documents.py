import json
from pathlib import Path

from pypdf import PdfReader
from docx import Document

from engine.config import ALLOWED_PATHS, ROOT_DIR, SUPPORTED_EXTENSIONS
from engine.security import SecurityError, validate_file_path

INDEX_PATH = ROOT_DIR / "document_index.json"
_index: list[dict] = []


def _extract_text(file_path: Path, snippet_only: bool = False) -> str:
    suffix = file_path.suffix.lower()
    try:
        if suffix == ".pdf":
            reader = PdfReader(str(file_path))
            pages = reader.pages[:1] if snippet_only else reader.pages
            return "\n".join(page.extract_text() or "" for page in pages)
        if suffix == ".docx":
            doc = Document(str(file_path))
            paragraphs = doc.paragraphs[:20] if snippet_only else doc.paragraphs
            return "\n".join(p.text for p in paragraphs)
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        return text[:2000] if snippet_only else text
    except Exception as e:
        return f"[Could not read file: {e}]"


def build_index() -> int:
    global _index
    _index = []
    for root in ALLOWED_PATHS:
        if not root.exists():
            continue
        for file_path in root.rglob("*"):
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            try:
                validate_file_path(file_path)
            except SecurityError:
                continue
            _index.append({
                "path": str(file_path),
                "name": file_path.name,
                "snippet": "",
            })
    INDEX_PATH.write_text(json.dumps(_index, indent=2), encoding="utf-8")
    return len(_index)


def _load_index() -> list[dict]:
    global _index
    if _index:
        return _index
    if INDEX_PATH.exists():
        _index = json.loads(INDEX_PATH.read_text(encoding="utf-8"))
        return _index
    build_index()
    return _index


def search_documents(query: str, limit: int = 5) -> list[dict]:
    query_lower = query.lower()
    words = query_lower.split()
    results = []
    for entry in _load_index():
        name_lower = entry["name"].lower()
        name_score = sum(1 for w in words if w in name_lower)
        if name_score > 0:
            results.append({
                "path": entry["path"],
                "name": entry["name"],
                "snippet": entry.get("snippet", "")[:200],
                "score": name_score * 2,
            })
            continue
        if entry.get("snippet"):
            snippet_lower = entry["snippet"].lower()
            snippet_score = sum(1 for w in words if w in snippet_lower)
            if snippet_score > 0:
                results.append({
                    "path": entry["path"],
                    "name": entry["name"],
                    "snippet": entry["snippet"][:200],
                    "score": snippet_score,
                })
    results.sort(key=lambda x: x["score"], reverse=True)
    return results[:limit]


def get_index_count() -> int:
    return len(_load_index())


def read_document(path: str) -> str:
    file_path = validate_file_path(path)
    text = _extract_text(file_path, snippet_only=False)
    max_chars = 8000
    if len(text) > max_chars:
        return text[:max_chars] + f"\n\n[Truncated — file has {len(text)} characters total]"
    return text
